
# In[60]:

from database import Ubion_data
import sys
import xlrd
import csv
import datetime
from docx import Document
from docx.oxml.ns import nsdecls 
from docx.oxml import parse_xml 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Cm, Pt, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from collections import Counter

filePath = "client/public/files/"+sys.argv[2]
f = open(filePath, 'r')
lines = csv.reader(f)
#for문에 들어가기전에, 모든 인스턴스들을 관리할 리스트 생성
menu_list = list()
univ_list = list()
#for문에 들어가기전에, 첫 줄에서는 작업안하도록 매개변수 생성
i = 0
for line in lines:
    #첫줄에서는 작업안하도록 if문 걸어주기
    if i > 0:
        univ_list.append(Ubion_data(line[1], line[0], line[3], line[5], line[6], line[7],
                                    line[10], line[11], line[12],line[13],line[9]))
    #첫줄 따로 저장
    elif i == 0:
        menu_list.append(Ubion_data(line[1], line[0], line[3], line[5], line[6], line[7],
                                    line[10], line[11], line[12],line[13],line[9]))
    i += 1
f.close()

#---------------------------연월 설정------------------------------

today = datetime.date.today()
month_ago = today.replace(day = 1)
year = str(month_ago.year)
month = int(str(month_ago.month))
month_2 = month-2
month_1 = month-1
if month-2 == -1:
    month_2 = 11
elif month-2 == 0:
    month_2 = 12
if month-1 == 0:
    month_1 = 12
    
if month_2 < 10:
    month_2_check = str(year)+'-0'+str(month_2)
else:
    month_2_check = str(year)+'-'+str(month_2)

if month_1 <10:
    month_1_check = str(year)+'-0'+str(month_1)
else:
    month_1_check = str(year)+'-'+str(month_1)

if month <10:
    month_check = str(year)+'-0'+str(month)
else:
    month_check = str(year)+'-'+str(month)

month_2_title = str(month_2)+'월'
month_1_title = str(month_1)+'월'
month_title = str(month)+'월'

#--------------------------------------1번------------------------------------------

issue_count_4 = 0
issue_end_4 = 0
issue_term_4 = 0
issue_count_5 = 0
issue_end_5 = 0
issue_term_5 = 0
issue_count_6 = 0
issue_end_6 = 0
issue_term_6 = 0
for a in univ_list:
    if month_2_check in a.request_date:
        issue_count_4 += 1
        if month_2_check in a.end_date:
            issue_term_4 += 1
    if a.progress == '완료' and month_2_check in a.end_date:
        issue_end_4 += 1
        
    if month_1_check in a.request_date:
        issue_count_5 += 1
        if month_1_check in a.end_date:
            issue_term_5 += 1
    if a.progress == '완료' and month_1_check in a.end_date:
        issue_end_5 += 1
        
    if month_check in a.request_date:
        issue_count_6 += 1
        if month_check in a.end_date:
            issue_term_6 += 1
    if a.progress == '완료' and month_check in a.end_date:
        issue_end_6 += 1
        
                
                
issue_count_total = issue_count_4+issue_count_5+issue_count_6
issue_end_total = issue_end_4+issue_end_5+issue_end_6
issue_term_total = issue_term_4+issue_term_5+issue_term_6
today = datetime.datetime.today().month

#--------------------------------------2번------------------------------------------

school_total= {}
school_2 = {}
school_1 = {}
school = {}
for lst in univ_list: 
    if lst.school not in school_total.keys(): 
        school_total[lst.school] = 1
        if month_2_check not in lst.request_date:
            if month_1_check not in lst.request_date:
                if month_check not in lst.request_date:
                    school_total[lst.school] = 0
    else: 
        school_total[lst.school] += 1
        if month_2_check not in lst.request_date:
            if month_1_check not in lst.request_date:
                if month_check not in lst.request_date:
                    school_total[lst.school] -= 1
        
    if lst.school not in school_2.keys(): 
        school_2[lst.school] = 1 
        if month_2_check not in lst.request_date:
            school_2[lst.school] = 0
    else:
        school_2[lst.school] += 1
        if month_2_check not in lst.request_date:
            school_2[lst.school] -= 1
            
    if lst.school not in school_1.keys(): 
        school_1[lst.school] = 1 
        if month_1_check not in lst.request_date:
            school_1[lst.school] = 0
    else:
        school_1[lst.school] += 1
        if month_1_check not in lst.request_date:
            school_1[lst.school] -= 1
            
    if lst.school not in school.keys(): 
        school[lst.school] = 1 
        if month_check not in lst.request_date:
            school[lst.school] = 0
    else:
        school[lst.school] += 1
        if month_check not in lst.request_date:
            school[lst.school] -= 1
            
del(school_total[''])
del(school_2[''])
del(school_1[''])
del(school[''])
issue_total = list(school_total.values())
issue_2 = list(school_2.values())
issue_1= list(school_1.values())
issue = list(school.values())
issue_name = list(school_total.keys())
customer_list = []
for ls in range(len(issue)):
    customer_list += [(issue_name[ls],issue_total[ls],issue_2[ls],issue_1[ls],issue[ls])]
customer = sorted(customer_list, key = lambda x : (-x[1]))

for j in range(len(customer)):
    for i in range(len(customer)):
        if customer[i][1] == 0:
            del customer[i]
            break
#--------------------------------------3번------------------------------------------
            
handling = {}
for lst in univ_list: 
    if lst.progress not in handling.keys(): 
        handling[lst.progress] = 1
        if month_2_check not in lst.request_date:
            handling[lst.progress] = 0
    else: 
        handling[lst.progress] += 1
        if month_2_check not in lst.request_date:
            handling[lst.progress] -= 1
            
    if lst.progress not in handling.keys(): 
        handling[lst.progress] = 1
        if month_1_check not in lst.request_date:
            handling[lst.progress] = 0
    else: 
        handling[lst.progress] += 1
        if month_1_check not in lst.request_date:
            handling[lst.progress] -= 1
            
    if lst.progress not in handling.keys(): 
        handling[lst.progress] = 1
        if month_check not in lst.request_date:
            handling[lst.progress] = 0
    else: 
        handling[lst.progress] += 1
        if month_check not in lst.request_date:
            handling[lst.progress] -= 1

handling = Counter(handling)
handling = dict(handling)
result = 0
for key,value in handling.items():
    result += value
    if key == '접수':
        acp = value
    elif key == '진행':
        prog = value
    elif key == '완료':
        comp = value
    elif key == '검토':
        rev = value
    elif key == '종결':
        end = value

acp_rate = round(acp/result*100,1)
prog_rate = round(prog/result*100,1)
comp_rate = round(comp/result*100,1)
rev_rate = round(rev/result*100,1)
end_rate = round(end/result*100,1)

acp_total_rate = round(acp_rate+prog_rate+comp_rate+rev_rate+end_rate,1)
if acp_total_rate < 100:
    result = 100 - acp_total_rate
    acp_rate += round(result,1)

#--------------------------------------4번------------------------------------------

product_check = {}
for lst in univ_list:
    if lst.product not in product_check.keys(): 
        product_check[lst.product] = 1
        if month_2_check not in lst.request_date:
            product_check[lst.product] = 0
    else: 
        product_check[lst.product] += 1
        if month_2_check not in lst.request_date:
            product_check[lst.product] -= 1
            
    if lst.product not in product_check.keys(): 
        product_check[lst.product] = 1
        if month_1_check not in lst.request_date:
            product_check[lst.product] = 0
    else: 
        product_check[lst.product] += 1
        if month_1_check not in lst.request_date:
            product_check[lst.product] -= 1
            
    if lst.product not in product_check.keys(): 
        product_check[lst.product] = 1
        if month_check not in lst.request_date:
            product_check[lst.product] = 0
    else: 
        product_check[lst.product] += 1
        if month_check not in lst.request_date:
            product_check[lst.product] -= 1
            
            
product_sum = Counter(product_check)
product_sum = dict(product_sum)
del(product_sum[''])
del(product_sum['공지사항'])
result = 0
for key,value in product_sum.items():
    result += value
    if key == 'LMS':
        lms = value
    if key == '기타':
        etc = value
    if key == 'LMS 외 제품':
        not_lms = value

lms_rate = round(lms/result*100,1)
etc_rate = round(etc/result*100,1)
not_lms_rate = round(not_lms/result*100,1)
lms_total_rate = round(lms_rate+etc_rate+not_lms_rate,1)
if lms_total_rate < 100:
    result = 100-lms_total_rate
    lms_rate += round(result,1)

#--------------------------------------5번------------------------------------------
group_check = {}
for lst in univ_list:
    if lst.group == '':
        lst.group.replace('','미분류')
    if lst.group not in group_check.keys(): 
        group_check[lst.group] = 1
        if month_2_check not in lst.request_date:
            group_check[lst.group] = 0
    else: 
        group_check[lst.group] += 1
        if month_2_check not in lst.request_date:
            group_check[lst.group] -= 1
            
    if lst.group not in group_check.keys(): 
        group_check[lst.group] = 1
        if month_1_check not in lst.request_date:
            group_check[lst.group] = 0
    else: 
        group_check[lst.group] += 1
        if month_1_check not in lst.request_date:
            group_check[lst.group] -= 1
            
    if lst.group not in group_check.keys(): 
        group_check[lst.group] = 1
        if month_check not in lst.request_date:
            group_check[lst.group] = 0
    else: 
        group_check[lst.group] += 1
        if month_check not in lst.request_date:
            group_check[lst.group] -= 1
            

group_sum = group_check
group_sum = dict(group_sum)
result = 0
for key,value in group_sum.items():
    result += value
    if key == '운영지원요청':
        sup_request = value
    elif key == '사용방법문의':
        how_to_use = value
    elif '오류' in key:
        error = value
    elif '수정' in key:
        improv = value
    elif key == '':
        unclass = value
sup_request_rate = round(sup_request/result*100,1)
how_to_use_rate = round(how_to_use/result*100,1)
error_rate = round(error/result*100,1)
improv_rate = round(improv/result*100,1)
unclass_rate = round(unclass/result*100,1)

result=0
sup_total_rate = round(sup_request_rate+how_to_use_rate+error_rate+improv_rate+unclass_rate,1)
if sup_total_rate < 100:
    result = 100-sup_total_rate
    sup_request_rate += round(result,1)
sup_request_rate = round(sup_request_rate,1)

#--------------------------------------6번------------------------------------------

disposer_check = {}
for lst in univ_list:
    if lst.disposer1 not in disposer_check.keys(): 
        disposer_check[lst.disposer1] = 1
        if month_2_check not in lst.request_date:
            disposer_check[lst.disposer1] = 0
    else: 
        disposer_check[lst.disposer1] += 1
        if month_2_check not in lst.request_date:
            disposer_check[lst.disposer1] -= 1
            
    if lst.disposer1 not in disposer_check.keys(): 
        disposer_check[lst.disposer1] = 1
        if month_1_check not in lst.request_date:
            disposer_check[lst.disposer1] = 0
    else: 
        disposer_check[lst.disposer1] += 1
        if month_1_check not in lst.request_date:
            disposer_check[lst.disposer1] -= 1
            
    if lst.disposer1 not in disposer_check.keys(): 
        disposer_check[lst.disposer1] = 1
        if month_check not in lst.request_date:
            disposer_check[lst.disposer1] = 0
    else: 
        disposer_check[lst.disposer1] += 1
        if month_check not in lst.request_date:
            disposer_check[lst.disposer1] -= 1
    
    if lst.disposer2 not in disposer_check.keys(): 
        disposer_check[lst.disposer2] = 1
        if month_2_check not in lst.request_date:
            disposer_check[lst.disposer2] = 0
    else: 
        disposer_check[lst.disposer2] += 1
        if month_2_check not in lst.request_date:
            disposer_check[lst.disposer2] -= 1
            
    if lst.disposer2 not in disposer_check.keys(): 
        disposer_check[lst.disposer2] = 1
        if month_1_check not in lst.request_date:
            disposer_check[lst.disposer2] = 0
    else: 
        disposer_check[lst.disposer2] += 1
        if month_1_check not in lst.request_date:
            disposer_check[lst.disposer2] -= 1
            
    if lst.disposer2 not in disposer_check.keys(): 
        disposer_check[lst.disposer2] = 1
        if month_check not in lst.request_date:
            disposer_check[lst.disposer2] = 0
    else: 
        disposer_check[lst.disposer2] += 1
        if month_check not in lst.request_date:
            disposer_check[lst.disposer2] -= 1
    
del(disposer_check[''])
disposers = sorted(disposer_check.items(), key =lambda x:(-x[1]))

k=0
for i in range(len(disposers)):
    for j in range(len(disposers)):
        if j>15:
            del disposers[j]
            break
            
#-----------------------------------7번----------------------------------------------------
blue_soft = {}
service_team = {}
for lst in univ_list:
    if lst.disposer1 == '블루소프트':
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_2_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_2_check not in lst.request_date:
                blue_soft[lst.group] -= 1
                
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_1_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_1_check not in lst.request_date:
                blue_soft[lst.group] -= 1
                
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_check not in lst.request_date:
                blue_soft[lst.group] -= 1
                
    if lst.disposer2 == '블루소프트':
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_2_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_2_check not in lst.request_date:
                blue_soft[lst.group] -= 1
                
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_1_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_1_check not in lst.request_date:
                blue_soft[lst.group] -= 1
                
        if lst.group not in blue_soft.keys(): 
            blue_soft[lst.group] = 1
            if month_check not in lst.request_date:
                blue_soft[lst.group] = 0
        else: 
            blue_soft[lst.group] += 1
            if month_check not in lst.request_date:
                blue_soft[lst.group] -= 1

for lst in univ_list:
    if len(lst.disposer1) == 3:
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_2_check not in lst.request_date:
                service_team[lst.group] = 0
        else: 
            service_team[lst.group] += 1
            if month_2_check not in lst.request_date:
                service_team[lst.group] -= 1
                
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_1_check not in lst.request_date:
                service_team[lst.group] = 0
        else: 
            service_team[lst.group] += 1
            if month_1_check not in lst.request_date:
                service_team[lst.group] -= 1
                
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_check not in lst.request_date:
                service_team[lst.group] = 0
        else: 
            service_team[lst.group] += 1
            if month_check not in lst.request_date:
                service_team[lst.group] -= 1
                
    if len(lst.disposer2) == 3:
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_2_check not in lst.request_date:
                service_team[lst.group] = 0
        else:
            service_team[lst.group] += 1
            if month_2_check not in lst.request_date:
                service_team[lst.group] -= 1
                
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_1_check not in lst.request_date:
                service_team[lst.group] = 0
        else: 
            service_team[lst.group] += 1
            if month_1_check not in lst.request_date:
                service_team[lst.group] -= 1
                
        if lst.group not in service_team.keys(): 
            service_team[lst.group] = 1
            if month_check not in lst.request_date:
                service_team[lst.group] = 0
        else: 
            service_team[lst.group] += 1
            if month_check not in lst.request_date:
                service_team[lst.group] -= 1
blue_soft = dict(blue_soft)
service_team = dict(service_team)

result = 0
for key,value in blue_soft.items():
    result += value
    if key == '운영지원요청':
        blue_soft_sup_request = value
    elif key == '사용방법문의':
        blue_soft_how_to_use = value
    elif '오류' in key:
        blue_soft_error = value
    elif '수정' in key:
        blue_soft_improv = value
    elif key == '':
        blue_soft_unclass = value
blue_soft_sup_request_rate = round(blue_soft_sup_request/result*100,1)
blue_soft_how_to_use_rate = round(blue_soft_how_to_use/result*100,1)
blue_soft_error_rate = round(blue_soft_error/result*100,1)
blue_soft_improv_rate = round(blue_soft_improv/result*100,1)
blue_soft_unclass_rate = round(blue_soft_unclass/result*100,1)
blue_soft_total_rate = round(blue_soft_sup_request_rate+blue_soft_how_to_use_rate+
                        blue_soft_error_rate+blue_soft_improv_rate+blue_soft_unclass_rate,1)
if blue_soft_total_rate < 100:
    result = 100 - blue_soft_total_rate
    blue_soft_sup_request_rate += round(result,1)
    
result = 0
for key,value in service_team.items():
    result += value
    if key == '운영지원요청':
        service_team_sup_request = value
    elif key == '사용방법문의':
        service_team_how_to_use = value
    elif '오류' in key:
        service_team_error = value
    elif '수정' in key:
        service_team_improv = value
    elif key == '':
        service_team_unclass = value
service_team_sup_request_rate = round(service_team_sup_request/result*100,1)
service_team_how_to_use_rate = round(service_team_how_to_use/result*100,1)
service_team_error_rate = round(service_team_error/result*100,1)
service_team_improv_rate = round(service_team_improv/result*100,1)
service_team_unclass_rate = round(service_team_unclass/result*100,1)
service_team_total_rate = round(service_team_sup_request_rate+service_team_how_to_use_rate+
                        service_team_error_rate+service_team_improv_rate+service_team_unclass_rate,1)
if service_team_total_rate < 100:
    result = 100 - service_team_total_rate
    service_team_sup_request_rate += round(result,1)
#-----------------------------------8번----------------------------------------------------

charge_check = {}

for lst in univ_list:
    if lst.charge[0:3] not in charge_check.keys(): 
        charge_check[lst.charge[0:3]] = 1
        if month_2_check not in lst.request_date:
            charge_check[lst.charge[0:3]] = 0
    else: 
        charge_check[lst.charge[0:3]] += 1
        if month_2_check not in lst.request_date:
            charge_check[lst.charge[0:3]] -= 1

    if lst.charge[0:3] not in charge_check.keys(): 
        charge_check[lst.charge[0:3]] = 1
        if month_1_check not in lst.request_date:
            charge_check[lst.charge[0:3]] = 0
    else: 
        charge_check[lst.charge[0:3]] += 1
        if month_1_check not in lst.request_date:
            charge_check[lst.charge[0:3]] -= 1

    if lst.charge[0:3] not in charge_check.keys(): 
        charge_check[lst.charge[0:3]] = 1
        if month_check not in lst.request_date:
            charge_check[lst.charge[0:3]] = 0
    else: 
        charge_check[lst.charge[0:3]] += 1
        if month_check not in lst.request_date:
            charge_check[lst.charge[0:3]] -= 1
del(charge_check[''])
charges = sorted(charge_check.items(), key =lambda x:(-x[1]))
for i in range(len(charges)):
    for j in range(len(charges)):
        if j>5: #이거 숫자 바꾸면 상위 몇명 제외하고 리스트에서 삭제됨
            del charges[j]
            break

#-----------------------------------9번----------------------------------------------------

for a in univ_list:
    if '퀴즈' in a.keyword :
        a.keyword='퀴즈'
    elif '동영상' in a.keyword : 
        a.keyword='동영상'
    elif '온라인출석' in a.keyword : 
        a.keyword='동영상'
    elif '학습진도' in a.keyword : 
        a.keyword='동영상'
    elif '강의실 홈' in a.keyword : 
        a.keyword='강의실홈'
    elif '강의실홈' in a.keyword : 
        a.keyword='강의실홈'
    elif '사용자등록' in a.keyword : 
        a.keyword='사용자등록'
    elif '사용자일괄등록' in a.keyword : 
        a.keyword='사용자등록'  
    elif '사용자일괄추가' in a.keyword : 
        a.keyword='사용자등록' 
    elif '성적' in a.keyword : 
        a.keyword='성적'
    elif '비교과' in a.keyword : 
        a.keyword='비교과' 
    elif '자이닉스' in a.keyword : 
        a.keyword='자이닉스'
    elif '핵심역량' in a.keyword : 
        a.keyword='핵심역량' 
    elif '강의계획서' in a.keyword : 
        a.keyword='강의계획서'
    elif '수강' in a.keyword : 
        a.keyword='수강' 
    elif '개인정보' in a.keyword : 
        a.keyword='개인정보'     

keyword = {}
for a in univ_list:
    if a.keyword not in keyword.keys(): 
        keyword[a.keyword] = 1
        if month_2_check not in a.request_date:
            keyword[a.keyword] = 0
    else:
        keyword[a.keyword] += 1
        if month_2_check not in a.request_date:
            keyword[a.keyword] -= 1
    
    if a.keyword not in keyword.keys(): 
        keyword[a.keyword] = 1
        if month_1_check not in a.request_date:
            keyword[a.keyword] = 0
    else:
        keyword[a.keyword] += 1
        if month_1_check not in a.request_date:
            keyword[a.keyword] -= 1
            
    if a.keyword not in keyword.keys(): 
        keyword[a.keyword] = 1
        if month_check not in a.request_date:
            keyword[a.keyword] = 0
    else:
        keyword[a.keyword] += 1
        if month_check not in a.request_date:
            keyword[a.keyword] -= 1
#빈칸제거

del(keyword[''])
keyword = sorted(keyword.items(), key =lambda x:(-x[1]))
for i in range(len(keyword)):
    for j in range(len(keyword)):
        if j>6: #이거 숫자 바꾸면 상위 몇명 제외하고 리스트에서 삭제됨
            del keyword[j]
            break
#-----------------------------------WORD생성----------------------------------------------------
#당일 날짜로 저장을위한 변수
savefile = str(datetime.date.today())

#문서생성
document = Document()
#페이지 방향 설정
current_section = document.sections[-1]
current_section.page_height = Mm(360)
current_section.page_width = Mm(270)
new_width, new_height = current_section.page_height, current_section.page_width
current_section.orientation = WD_ORIENT.LANDSCAPE
current_section.page_width = new_width
current_section.page_height = new_height

#글꼴 설정
style = document.styles['Normal']
style.font.name = '나눔바른고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔바른고딕')
#style.font.size = Pt(30)

#용지 설정
sections = document.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

hdstyle = document.add_paragraph('\n\n\n\n\n데이터 수집 범위')
hdstyle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#새 페이지 생성
document.add_page_break()

#-----------------------------------1번----------------------------------------------------

#헤더 생성
service = document.add_paragraph('1.서비스운영 이슈현황')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
table = document.add_table(rows=1, cols=5)
table.style = document.styles['Table Grid']
#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
hdr_cells[1].text = str('전체(분기)')
hdr_cells[2].text = str(month_2_title)
hdr_cells[3].text = str(month_1_title)
hdr_cells[4].text = str(month_title)
#테이블 텍스트 중앙정렬
for i in range(0,5):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
#테이블 행 추가
for i in range(0,3):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('전체이슈')
        row_cells[1].text = str(issue_count_total)
        row_cells[2].text = str(issue_count_4)
        row_cells[3].text = str(issue_count_5)
        row_cells[4].text = str(issue_count_6)
    elif i == 1:
        row_cells[0].text = str('완료처리\n('+savefile+'기준)')
        row_cells[1].text = str(issue_end_total)
        row_cells[2].text = str(issue_end_4)
        row_cells[3].text = str(issue_end_5)
        row_cells[4].text = str(issue_end_6)
    else:
        row_cells[0].text = str('해당기간에 완료처리한 이슈')
        row_cells[1].text = str(issue_term_total)
        row_cells[2].text = str(issue_term_4)
        row_cells[3].text = str(issue_term_5)
        row_cells[4].text = str(issue_term_6)
    #테이블 텍스트 중앙정렬
    for j in range(0,5):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(3)
#새 페이지 생성
document.add_page_break()

#-----------------------------------2번----------------------------------------------------

service = document.add_paragraph('2.고객사별 이슈')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#글꼴 설정
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
    
information = document.add_paragraph('정렬순서 : 전체 이슈수(내림차순)')
information.aligment = WD_ALIGN_PARAGRAPH.LEFT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(12)

table = document.add_table(rows=1, cols=5)
table.style = document.styles['Table Grid']
#테이블 생성
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('학교')
hdr_cells[1].text = str('전체이슈')
hdr_cells[2].text = str(month_2_title)
hdr_cells[3].text = str(month_1_title)
hdr_cells[4].text = str(month_title)
for i in range(0,5):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
#테이블 행 추가
for i in range(len(customer)):
    row_cells = table.add_row().cells
    for j in range(0,5):
        row_cells[j].text = str(customer[i][j])
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for k in range(1,5):
        row_cells[k].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(1)
#가장 큰 수 색 넣기
#각각 따로 설정해야함
total_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
one_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
two_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
three_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
four_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
five_color = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))

#전체이슈
j = 0
for i in range(len(customer)):
    result = int(customer[i][1])
    if j < result:
        j = result
j = str(j)
for i in range(len(customer)):
    if str(customer[i][1]) == j:
        if i == 0:
            table.rows[i+1].cells[1]._tc.get_or_add_tcPr().append(total_color)
#n-2월
j = 0
for i in range(len(customer)):
    result = int(customer[i][2])
    if j < result:
        j = result
j = str(j)
for i in range(len(customer)):
    if str(customer[i][2]) == j:
        table.rows[i+1].cells[2]._tc.get_or_add_tcPr().append(one_color)
#n-1월 
j = 0
for i in range(len(customer)):
    result = int(customer[i][3])
    if j < result:
        j = result
j = str(j)
for i in range(len(customer)):
    if str(customer[i][3]) == j:
        table.rows[i+1].cells[3]._tc.get_or_add_tcPr().append(two_color)
#n월
j = 0
for i in range(len(customer)):
    result = int(customer[i][4])
    if j < result:
        j = result
j = str(j)
for i in range(len(customer)):
    if str(customer[i][4]) == j:
        table.rows[i+1].cells[4]._tc.get_or_add_tcPr().append(three_color)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(1)
        
document.add_page_break()
        
#-----------------------------------3번----------------------------------------------------
    
service = document.add_paragraph('3.이슈처리 상황')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
table = document.add_table(rows=1, cols=6)
table.style = document.styles['Table Grid']

#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
hdr_cells[1].text = str('접수')
hdr_cells[2].text = str('진행')
hdr_cells[3].text = str('완료')
hdr_cells[4].text = str('검토')
hdr_cells[5].text = str('종결')
#테이블 텍스트 중앙정렬
for i in range(0,6):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,2):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('전체이슈')
        row_cells[1].text = str(acp)
        row_cells[2].text = str(prog)
        row_cells[3].text = str(comp)
        row_cells[4].text = str(rev)
        row_cells[5].text = str(end)
    else:
        row_cells[0].text = str('비율(%)')
        row_cells[1].text = str(round(acp_rate,1))
        row_cells[2].text = str(round(prog_rate,1))
        row_cells[3].text = str(round(comp_rate,1))
        row_cells[4].text = str(round(rev_rate,1))
        row_cells[5].text = str(round(end_rate,1))
    #테이블 텍스트 중앙정렬
    for j in range(0,6):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정         
for row in table.rows:
    row.height = Cm(3)
    
document.add_page_break()
        
#-----------------------------------4번----------------------------------------------------
    
service = document.add_paragraph('4.제품별 이슈')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
table = document.add_table(rows=1, cols=4)
table.style = document.styles['Table Grid']

#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
hdr_cells[1].text = str('LMS')
hdr_cells[2].text = str('LMS 외 제품')
hdr_cells[3].text = str('기타(비제품영역)')
#테이블 텍스트 중앙정렬
for i in range(0,4):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,2):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('이슈수')
        row_cells[1].text = str(lms)
        row_cells[2].text = str(not_lms)
        row_cells[3].text = str(etc)
    else:
        row_cells[0].text = str('비율(%)')
        row_cells[1].text = str(lms_rate)
        row_cells[2].text = str(not_lms_rate)
        row_cells[3].text = str(etc_rate)
    #테이블 텍스트 중앙정렬
    for j in range(0,4):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정         
for row in table.rows:
    row.height = Cm(3)

document.add_page_break()
        
#-----------------------------------5번----------------------------------------------------
    
service = document.add_paragraph('5.문의내용별 분류')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
table = document.add_table(rows=1, cols=6)
table.style = document.styles['Table Grid']

#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
hdr_cells[1].text = str('운영지원요청')
hdr_cells[2].text = str('사용방법문의')
hdr_cells[3].text = str('오류/장애')
hdr_cells[4].text = str('수정/개선요청')
hdr_cells[5].text = str('미분류')
#테이블 텍스트 중앙정렬
for i in range(0,6):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,2):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('이슈수')
        row_cells[1].text = str(sup_request)
        row_cells[2].text = str(how_to_use)
        row_cells[3].text = str(error)
        row_cells[4].text = str(improv)
        row_cells[5].text = str(unclass)
    else:
        row_cells[0].text = str('비율(%)')
        row_cells[1].text = str(sup_request_rate)
        row_cells[2].text = str(how_to_use_rate)
        row_cells[3].text = str(error_rate)
        row_cells[4].text = str(improv_rate)
        row_cells[5].text = str(unclass_rate)
    #테이블 텍스트 중앙정렬
    for j in range(0,6):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정         
for row in table.rows:
    row.height = Cm(3)
    
document.add_page_break()
        
#-----------------------------------6번----------------------------------------------------
    
service = document.add_paragraph('6.문의처리 참여자')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
table = document.add_table(rows=1, cols=9)
table.style = document.styles['Table Grid']

#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('처리자')
for i in range(0,8):
    hdr_cells[i+1].text = str(disposers[i][0])
#테이블 텍스트 중앙정렬
for i in range(0,9):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,3):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('이슈수')
        for j in range(0,8):
            row_cells[j+1].text = str(disposers[j][1])
    elif i == 1:
        row_cells[0].text = str('처리자')
        for j in range(0,8):
            row_cells[j+1].text = str(disposers[j+8][0])
    else:
        row_cells[0].text = str('이슈수')
        for j in range(0,8):
            row_cells[j+1].text = str(disposers[j+8][1])
    #테이블 텍스트 중앙정렬
    for j in range(0,9):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
disposers_line_1 = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
disposers_line_2 = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
disposers_line_3 = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
disposers_line_4 = parse_xml(r'<w:shd {} w:fill="73FDEA"/>'.format(nsdecls('w')))
table.rows[0].cells[0]._tc.get_or_add_tcPr().append(disposers_line_1)
table.rows[1].cells[0]._tc.get_or_add_tcPr().append(disposers_line_2)
table.rows[2].cells[0]._tc.get_or_add_tcPr().append(disposers_line_3)
table.rows[3].cells[0]._tc.get_or_add_tcPr().append(disposers_line_4)
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정         
for row in table.rows:
    row.height = Cm(3)
    
document.add_page_break()

#--------------------------------------7번-----------------------------------------------

#헤더 생성
service = document.add_paragraph('7.주요 참여자 이슈 대응 내용')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
    
#테이블 생성
service = document.add_paragraph()
table = document.add_table(rows=1, cols=6)
table.style = document.styles['Table Grid']
#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('서비스운영팀')
hdr_cells[1].text = str('운영지원요청')
hdr_cells[2].text = str('사용방법문의')
hdr_cells[3].text = str('오류/장애')
hdr_cells[4].text = str('수정/개선요청')
hdr_cells[5].text = str('미분류')
table.rows[0].cells[0]._tc.get_or_add_tcPr().append(four_color)

#테이블 텍스트 중앙정렬
for i in range(0,6):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,2):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('총 이슈건')
        row_cells[1].text = str(service_team_sup_request)
        row_cells[2].text = str(service_team_how_to_use)
        row_cells[3].text = str(service_team_error)
        row_cells[4].text = str(service_team_improv)
        row_cells[5].text = str(service_team_unclass)
    else:
        row_cells[0].text = str('비율(%)')
        row_cells[1].text = str(service_team_sup_request_rate)
        row_cells[2].text = str(service_team_how_to_use_rate)
        row_cells[3].text = str(service_team_error_rate)
        row_cells[4].text = str(service_team_improv_rate)
        row_cells[5].text = str(service_team_unclass_rate)
    #테이블 텍스트 중앙정렬
    for j in range(0,6):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(2)

    #서비스운영팀
document.add_paragraph()
#테이블 생성
table = document.add_table(rows=1, cols=6)
table.style = document.styles['Table Grid']
#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('블루소프트')
hdr_cells[1].text = str('운영지원요청')
hdr_cells[2].text = str('사용방법문의')
hdr_cells[3].text = str('오류/장애')
hdr_cells[4].text = str('수정/개선요청')
hdr_cells[5].text = str('미분류')
table.rows[0].cells[0]._tc.get_or_add_tcPr().append(five_color)

#테이블 텍스트 중앙정렬
for i in range(0,6):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for i in range(0,2):
    row_cells = table.add_row().cells
    if i == 0:
        row_cells[0].text = str('총 이슈건')
        row_cells[1].text = str(blue_soft_sup_request)
        row_cells[2].text = str(blue_soft_how_to_use)
        row_cells[3].text = str(blue_soft_error)
        row_cells[4].text = str(blue_soft_improv)
        row_cells[5].text = str(blue_soft_unclass)
    else:
        row_cells[0].text = str('비율(%)')
        row_cells[1].text = str(blue_soft_sup_request_rate)
        row_cells[2].text = str(blue_soft_how_to_use_rate)
        row_cells[3].text = str(blue_soft_error_rate)
        row_cells[4].text = str(blue_soft_improv_rate)
        row_cells[5].text = str(blue_soft_unclass_rate)
    #테이블 텍스트 중앙정렬
    for j in range(0,6):
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(i+1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(2)
    
document.add_page_break()

#---------------------------------8번---------------------------------------------------

#헤더 생성
service = document.add_paragraph('8.이슈 담당자')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT

for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
#테이블 생성
document.add_paragraph()
table = document.add_table(rows=1, cols=6)
table.style = document.styles['Table Grid']
#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
for i in range(0,5):
    hdr_cells[i+1].text = str(charges[i][0])

#테이블 텍스트 중앙정렬
for i in range(0,6):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


row_cells = table.add_row().cells
row_cells[0].text = str('이슈수')
for i in range(0,5):
    row_cells[i+1].text = str(charges[i][1])

#테이블 텍스트 중앙정렬
for j in range(0,6):
    row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(3)
        
document.add_page_break()

#-------------------------------9번-------------------------------------------

#헤더 생성
service = document.add_paragraph('9.LMS이슈(주요)키워드 분석')
service.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
for run in paragraph.runs:
    run.font.size = Pt(30)
    
#테이블 생성
document.add_paragraph()
table = document.add_table(rows=1, cols=7)
table.style = document.styles['Table Grid']
#테이블에 값 입력
hdr_cells = table.rows[0].cells
hdr_cells[0].text = str('')
for i in range(0,6):
    hdr_cells[i+1].text = str(keyword[i][0])

#테이블 텍스트 중앙정렬
for i in range(0,7):
    hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

row_cells = table.add_row().cells
row_cells[0].text = str('이슈수')
for i in range(0,6):
    row_cells[i+1].text = str(keyword[i][1])
    
    #테이블 텍스트 중앙정렬
for j in range(0,7):
    row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
#테이블 글꼴 설정
for row in table.rows: 
    for cell in row.cells: 
        paragraphs = cell.paragraphs 
        for paragraph in paragraphs: 
            for run in paragraph.runs:
                font = run.font 
                font.size= Pt(15)
#셀 높이 설정
for row in table.rows:
    row.height = Cm(3)
#docx파일 생성을 위한 save('파일명')
fileName = sys.argv[2]
ffile = fileName.split(".")
document.save("client/public/covtFiles/"+ffile[0]+".docx")

print('함수종료')
